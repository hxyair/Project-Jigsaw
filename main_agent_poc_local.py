# ==============================================================================
# main_agent_poc_local.py
#
# Final Version - Uses GLM-4.5-Air for ALL agents (Specialists + PI)
# - Addresses timeout issues by using a faster, more reliable model.
# - Includes asyncio.gather for concurrent specialist calls.
# - Includes generalized prompts and optimized save_to_word.
# - Includes timing logs.
# - All comments are in English.
# ==============================================================================
import os
import re
import asyncio # Import asyncio for concurrency
import time  # <--- IMPORT TIME MODULE (Fix for NameError)
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP
from zhipuai import ZhipuAI

# --- Load Environment Variables ---
load_dotenv()
ZHIPU_API_KEY = os.getenv("ZHIPU_API_KEY")
if not ZHIPU_API_KEY:
    print("⚠️ ZHIPU_API_KEY not found...")
    exit(1)

# --- Define Models & Initialize Client ---
# NOW USING AIR MODEL FOR EVERYTHING FOR SPEED AND RELIABILITY
MODEL_TO_USE = "glm-4.5-air"
API_TIMEOUT = 300.0 # 5 minutes timeout for individual API calls

try:
    client = ZhipuAI(api_key=ZHIPU_API_KEY, timeout=API_TIMEOUT) # Pass timeout to client
    print(f"✅ Zhipu AI client initialized (API Timeout: {API_TIMEOUT}s).")
    print(f"✅ ALL Agents will use: {MODEL_TO_USE}")
except Exception as e:
    print(f"⚠️ Zhipu AI client initialization failed: {e}")
    exit(1)

# --- Helper: Save Report (Optimized, No Change) ---
def save_to_word(idea: str, content: str) -> str:
    raw_title = "_".join(idea.split()[:4])
    short_title = re.sub(r'[\\/*?:"<>|]', "", raw_title)
    short_title = short_title if short_title else "Untitled_Project" # Ensure not empty
    date_suffix = datetime.now().strftime("%m%d")
    filename = f"{short_title}-{date_suffix}.docx"
    os.makedirs("reports", exist_ok=True)
    filepath = os.path.join("reports", filename)
    doc = Document()
    doc.add_heading(f"R&D Project Proposal: {idea}", level=1)
    doc.add_paragraph(content)
    try:
        doc.save(filepath)
        print(f"✅ Word report saved (sanitized): {filepath}")
    except Exception as e:
        print(f"❌ Error saving Word document: {e}")
        try: # Fallback save
            fallback_filename = f"Fallback_Report_{date_suffix}.docx"
            filepath = os.path.join("reports", fallback_filename)
            doc.save(filepath)
            print(f"⚠️ Saved with fallback name due to error: {filepath}")
        except Exception as fallback_e:
            print(f"❌ Fallback save also failed: {fallback_e}")
            return "Error: File could not be saved."
    return os.path.abspath(filepath)

# --- Helper: Call Zhipu AI (No Change, includes robustness checks) ---
def call_glm_model(prompt: str, model_name: str) -> str:
    """Reusable helper to call a specified Zhipu AI model."""
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
        )
        # Check if response or choices might be None or empty
        if not response or not response.choices:
             print(f"❌ Empty response received from Zhipu AI ({model_name}).")
             return f"Error: Empty response received from {model_name}."
        # Safely access the message content
        message = response.choices[0].message
        content = getattr(message, 'content', None)
        if content is None:
             print(f"❌ Response message content is missing from Zhipu AI ({model_name}). Response: {response}")
             return f"Error: Response message content missing from {model_name}."
        return content.strip()
    except Exception as e:
        print(f"❌ Error calling Zhipu AI ({model_name}): {type(e).__name__} - {e}")
        return f"Error: Failed to get response from {model_name}. Details: {str(e)}"

# --- MCP Server & Agent Definitions ---
mcp = FastMCP("Collaborative R&D Planning Server")

# --- Specialist Agent Definitions (Now using MODEL_TO_USE) ---
# --- Using shortened prompts for brevity, replace with your full generalized prompts ---

@mcp.tool(description="Conduct background research...")
def background_agent(idea: str) -> str:
    prompt = f"""
    Conduct detailed background research for this project idea: "{idea}"
    Focus on: Problem Statement, Alignment with Strategic Goals (use placeholders like [Relevant National Strategy], [Your Institution Name]), Market Gap.
    """
    print(f"...background_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...background_agent finished call.")
    return result

@mcp.tool(description="Describe technical framework...")
def technical_agent(idea: str) -> str:
    prompt = f"""
    Describe the technical framework for a project based on this idea: "{idea}"
    Focus on: Core Technology (e.g., AI Adapter), Architecture (e.g., Agent-Tool, MCP), Key Innovations, Rationale (vs alternatives, why specific tech choices).
    """
    print(f"...technical_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...technical_agent finished call.")
    return result

@mcp.tool(description="Analyze market potential...")
def market_agent(idea: str) -> str:
    prompt = f"""
    Conduct a market and competitor analysis for this project idea: "{idea}"
    Focus on: Target Applications/Industries (relevant to Industry 4.0/SMEs), Competitive Landscape (using Manpower/Materials/Method framework), Commercialization Potential (realistic revenue models/partners).
    """
    print(f"...market_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...market_agent finished call.")
    return result

@mcp.tool(description="Estimate budget...")
def budget_agent(idea: str) -> str:
    prompt = f"""
    Estimate a **generic one-year budget** for "{idea}", suitable for a seed R&D funding (range **$50k-$100k USD**).
    Break down into EOM (student/staff), Equipment (prototyping hardware), OOE (consumables, cloud/API, contingency 10-15%). Use placeholders like [Stipend Rate]. Present clearly (e.g., table).
    """
    print(f"...budget_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...budget_agent finished call.")
    return result

@mcp.tool(description="Propose timeline, KPIs...")
def planner_agent(idea: str) -> str:
    prompt = f"""
    Propose a **generic one-year project plan** for "{idea}" (seed phase).
    Include: Timeline & Milestones (e.g., Q1-Q4, key phases like Design, Develop, Test, Document), Example KPIs (grouped: Tech Advancement, Knowledge Creation, Talent Dev, Collaboration - state specifics depend on funding body), Risk Assessment (Tech, Data, Ethical, Model Behavior, Project Mgmt + mitigations).
    """
    print(f"...planner_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...planner_agent finished call.")
    return result

@mcp.tool(description="Assess impact and ESG...")
def impact_agent(idea: str) -> str:
    prompt = f"""
    Assess the broader impacts and significance of "{idea}"
    Focus on: Technological Impact, Societal Relevance (industry benefits, alignment with goals), Institutional Benefit (use placeholder [Your Institution Name]), ESG Considerations (Environmental, Social, Governance). Use formal tone.
    """
    print(f"...impact_agent starting call using {MODEL_TO_USE}...")
    result = call_glm_model(prompt, MODEL_TO_USE)
    print("...impact_agent finished call.")
    return result

# --------------------------------------------------------------------------
# PI Agent (Coordinator): Uses asyncio for specialists, synthesizes with AIR model
# --------------------------------------------------------------------------
@mcp.tool(description="Synthesize all sections into a generic R&D proposal using GLM-4-Air.")
async def pi_agent(idea: str): # Still async
    """Integrates specialist outputs concurrently, then synthesizes."""
    start_time = time.time() # Start timer
    print(f"[{datetime.now()}] 🧠 [PI Agent] Coordinating... Starting concurrent calls to 6 specialists using {MODEL_TO_USE}.")

    # --- Step 1: Run all 6 specialist agents CONCURRENTLY ---
    specialist_start_time = time.time()
    try:
        results = await asyncio.gather(
             asyncio.to_thread(background_agent, idea),
             asyncio.to_thread(technical_agent, idea),
             asyncio.to_thread(market_agent, idea),
             asyncio.to_thread(budget_agent, idea),
             asyncio.to_thread(planner_agent, idea),
             asyncio.to_thread(impact_agent, idea),
             return_exceptions=True
         )
        specialist_end_time = time.time()
        print(f"[{datetime.now()}] 🧠 [PI Agent] All specialist tasks completed concurrently in {specialist_end_time - specialist_start_time:.2f} seconds.")

        # Unpack results and check for exceptions
        specialist_outputs = {}
        agent_names = ['background', 'technical', 'market', 'budget', 'planner', 'impact']
        errors_found = False
        for i, result in enumerate(results):
            agent_name = agent_names[i]
            if isinstance(result, Exception):
                print(f"⚠️ Error received from {agent_name}_agent during concurrent execution: {result}")
                specialist_outputs[agent_name] = f"Error: Failed during execution - {result}"
                errors_found = True
            elif isinstance(result, str) and result.startswith("Error:"): # Check for error messages
                 print(f"⚠️ Error reported by {agent_name}_agent (API call failed or empty response).")
                 specialist_outputs[agent_name] = result
                 errors_found = True
            else:
                specialist_outputs[agent_name] = result # Store successful result
                print(f"✅ Received result from {agent_name}_agent.")

    except Exception as gather_err:
        specialist_end_time = time.time()
        print(f"[{datetime.now()}] ❌ Error during asyncio.gather for specialists after {specialist_end_time - specialist_start_time:.2f} seconds: {gather_err}")
        return {"status": "error", "message": f"Critical error during specialist agent execution: {gather_err}"}

    # --- Step 2: Create synthesis prompt ---
    synthesis_prompt = f"""
    You are a Principal Investigator (PI) writing a **generic one-year seed R&D project proposal**.
    The core project idea is: "{idea}"
    You have received draft sections from 6 specialist agents. Synthesize these into a formal, coherent, compelling proposal document. If inputs contain errors, note this.

    **Proposal Structure & Content Guidelines:**
    1.  **Title Page Placeholder:** "[Project Title Page: Title, PI, Institution, Date]".
    2.  **Executive Summary:** Concise overview (problem, solution, innovation, impact, resources).
    3.  **1. Project Background:** Combine problem, market gap, strategic alignment (use placeholders `[Relevant Strategy]`, `[Your Institution Name]'s Goals`).
    4.  **2. Project Objectives/Purpose:** State seed phase goals (Tech Innovation, Knowledge/Talent Dev, Collaboration).
    5.  **3. Proposed Approach / Technical Framework:** Core tech, architecture (Agent-Tool, MCP), innovations, rationale.
    6.  **4. Market Context & Competitor Analysis:** Applications/industries, competitive advantages (Manpower/Materials/Method).
    7.  **5. Project Plan:** 5.1 Timeline & Milestones. 5.2 Example KPIs (Tech, Knowledge, Talent, Collab). State "*Specific KPIs TBD*". 5.3 Risk Assessment & Mitigations.
    8.  **6. Budget Outline:** Estimated budget ($50k-$100k range) (EOM, Equipment, OOE). State "*Estimates only. Use placeholder [Detailed Budget Table...]*".
    9.  **7. Project Team Placeholder:** "**7. Project Team:** *[Details of PI, Co-PI(s), members, roles, expertise from [Relevant Dept 1], [Relevant Dept 2] here.]*"
    10. **8. Impact and Significance:** Tech impact, societal relevance, institutional benefits (use `[Your Institution Name]`), ESG considerations.
    11. **9. Transferable Assets & Future Work:** List potential assets, outline next steps.
    12. **10. Conclusion:** Summarize value proposition.
    13. **Tone:** Formal, professional, realistic for seed funding. Use placeholders `[...]` for specifics.

    --- [BEGIN SPECIALIST INPUTS] ---

    --- 1. Background Research ---
    {specialist_outputs.get('background', 'No content received or error occurred.')}

    --- 2. Technical Framework ---
    {specialist_outputs.get('technical', 'No content received or error occurred.')}

    --- 3. Market & Competitor Analysis ---
    {specialist_outputs.get('market', 'No content received or error occurred.')}

    --- 4. Budget (Seed Format) ---
    {specialist_outputs.get('budget', 'No content received or error occurred.')}

    --- 5. Timeline, Milestones, KPIs, Risks ---
    {specialist_outputs.get('planner', 'No content received or error occurred.')}

    --- 6. Impact & Significance ---
    {specialist_outputs.get('impact', 'No content received or error occurred.')}

    --- [END SPECIALIST INPUTS] ---

    Now, generate the complete, synthesized, generic R&D proposal adhering strictly to the guidelines.
    """

    # --- Step 3: Call MODEL_TO_USE (GLM-4.5-Air) for synthesis ---
    print(f"🧠 [PI Agent] Now calling {MODEL_TO_USE} for final synthesis...")
    synthesis_start_time = time.time()
    final_proposal_content = call_glm_model(synthesis_prompt, MODEL_TO_USE)
    synthesis_end_time = time.time()

    if final_proposal_content.startswith("Error:"):
        print(f"[{datetime.now()}] ❌ Error calling {MODEL_TO_USE} for synthesis after {synthesis_end_time - synthesis_start_time:.2f} seconds.")
        return {"status": "error", "message": f"Failed to get synthesis response from {MODEL_TO_USE}. Details: {final_proposal_content}"}
    else:
        print(f"[{datetime.now()}] 🧠 [PI Agent] Final generic proposal synthesized by {MODEL_TO_USE} in {synthesis_end_time - synthesis_start_time:.2f} seconds.")

    # --- Step 4: Save the final proposal ---
    save_start_time = time.time()
    filepath = save_to_word(idea, final_proposal_content)
    save_end_time = time.time()
    print(f"[{datetime.now()}] 💾 File saved in {save_end_time - save_start_time:.2f} seconds.")

    if "Error: File could not be saved." in filepath:
         return {"status": "error", "message": "Synthesis complete, but failed to save the final Word document."}

    end_time = time.time()
    print(f"[{datetime.now()}] ✅ PI Agent finished successfully in {end_time - start_time:.2f} seconds total.")

    final_status = "success" if not errors_found else "partial_success"
    final_message = f"✅ Generic proposal synthesized by {MODEL_TO_USE} and saved: {filepath}"
    if errors_found:
        final_message = f"⚠️ Proposal synthesized with potential gaps due to specialist errors. Saved: {filepath}"

    return {
        "status": final_status,
        "message": final_message,
        "file_path": filepath,
    }

# ==============================================================================
# 6. Run MCP Server
# ==============================================================================
if __name__ == "__main__":
    print("🚀 Starting Collaborative MCP Server...")
    mcp.run(transport="streamable-http")